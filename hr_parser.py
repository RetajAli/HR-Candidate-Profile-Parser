import json
import os
from io import BytesIO
from typing import List

import fitz
import pandas as pd
import streamlit as st
from docx import Document
from dotenv import load_dotenv
from huggingface_hub import InferenceClient
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate
from pydantic import BaseModel, Field


# --------------------------------------------------
# Application setup
# --------------------------------------------------

load_dotenv()

st.set_page_config(
    page_title="HR Candidate Profile Parser",
    page_icon="👤",
    layout="wide"
)


# --------------------------------------------------
# Candidate data structure
# --------------------------------------------------

class Education(BaseModel):
    degree: str = Field(
        default="",
        description="Candidate's degree"
    )
    institution: str = Field(
        default="",
        description="University or educational institution"
    )
    year: str = Field(
        default="",
        description="Graduation year"
    )


class Experience(BaseModel):
    role: str = Field(
        default="",
        description="Candidate's job role"
    )
    company: str = Field(
        default="",
        description="Company name"
    )
    years: str = Field(
        default="",
        description="Employment years or duration"
    )


class CandidateProfile(BaseModel):
    full_name: str = Field(
        default="",
        description="Candidate's full name"
    )
    email: str = Field(
        default="",
        description="Candidate's email address"
    )
    education: List[Education] = Field(
        default_factory=list,
        description="Candidate's education history"
    )
    skills: List[str] = Field(
        default_factory=list,
        description="Candidate's skills"
    )
    experience: List[Experience] = Field(
        default_factory=list,
        description="Candidate's work experience"
    )


# --------------------------------------------------
# File text extraction
# --------------------------------------------------

def extract_pdf_text(uploaded_file) -> str:
    """Extract text from an uploaded PDF file."""

    pdf_bytes = uploaded_file.getvalue()

    document = fitz.open(
        stream=pdf_bytes,
        filetype="pdf"
    )

    pages = []

    for page in document:
        page_text = page.get_text(
            "text",
            sort=True
        )

        if page_text.strip():
            pages.append(page_text)

    document.close()

    return "\n".join(pages).strip()


def extract_docx_text(uploaded_file) -> str:
    """Extract text from an uploaded DOCX file."""

    file_bytes = BytesIO(
        uploaded_file.getvalue()
    )

    document = Document(file_bytes)

    text_parts = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            text_parts.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            row_values = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    row_values.append(cell_text)

            if row_values:
                text_parts.append(
                    " | ".join(row_values)
                )

    return "\n".join(text_parts).strip()


def extract_txt_text(uploaded_file) -> str:
    """Extract text from an uploaded TXT file."""

    return uploaded_file.getvalue().decode(
        "utf-8",
        errors="ignore"
    ).strip()


def extract_resume_text(uploaded_file) -> str:
    """Choose the correct extraction method."""

    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        return extract_pdf_text(uploaded_file)

    if file_name.endswith(".docx"):
        return extract_docx_text(uploaded_file)

    if file_name.endswith(".txt"):
        return extract_txt_text(uploaded_file)

    raise ValueError(
        "Only PDF, DOCX and TXT files are supported."
    )


# --------------------------------------------------
# LangChain output parser
# --------------------------------------------------

def create_output_parser() -> JsonOutputParser:
    """Create a LangChain JSON output parser."""

    return JsonOutputParser(
        pydantic_object=CandidateProfile
    )


def create_prompt(resume_text: str) -> str:
    """Create the prompt sent to the AI model."""

    parser = create_output_parser()

    format_instructions = (
        parser.get_format_instructions()
    )

    prompt_template = PromptTemplate(
        template="""
You are an HR Candidate Profile Parser.

Read the resume text and extract the candidate information.

Follow these rules carefully:

1. Use only information found in the resume.
2. Do not invent or assume information.
3. Return the candidate's full name.
4. Return the candidate's email.
5. Return education as a list of objects.
6. Every education object must contain:
   - degree
   - institution
   - year
7. Return skills as a list of individual strings.
8. Return experience as a list of objects.
9. Every experience object must contain:
   - role
   - company
   - years
10. Use an empty string when a text value is missing.
11. Use an empty list when a list is missing.
12. Return valid JSON only.
13. Do not write explanations before or after the JSON.

Resume text:

{resume_text}

JSON formatting instructions:

{format_instructions}
""",
        input_variables=[
            "resume_text"
        ],
        partial_variables={
            "format_instructions":
                format_instructions
        }
    )

    return prompt_template.format(
        resume_text=resume_text[:12000]
    )


# --------------------------------------------------
# Hugging Face AI parsing
# --------------------------------------------------

def parse_candidate_profile(
    resume_text: str
) -> dict:
    """Send the resume to Hugging Face and parse the result."""

    token = os.getenv("HF_TOKEN")

    if not token:
        raise ValueError(
            "HF_TOKEN was not found. "
            "Open the .env file and add: "
            "HF_TOKEN=your_token"
        )

    prompt = create_prompt(resume_text)

    client = InferenceClient(
        token=token
    )

    response = client.chat_completion(
        model="Qwen/Qwen2.5-7B-Instruct",
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an HR resume parser. "
                    "Return valid JSON only."
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        max_tokens=1200,
        temperature=0.1
    )

    output_text = (
        response.choices[0]
        .message
        .content
    )

    if not output_text:
        raise ValueError(
            "The AI model returned an empty response."
        )

    parser = create_output_parser()

    parsed_result = parser.parse(
        output_text
    )

    validated_profile = (
        CandidateProfile.model_validate(
            parsed_result
        )
    )

    return validated_profile.model_dump()


# --------------------------------------------------
# Profile calculations
# --------------------------------------------------

def calculate_completeness(
    profile: dict
) -> int:
    """Calculate how complete the profile is."""

    checks = [
        bool(profile.get("full_name")),
        bool(profile.get("email")),
        bool(profile.get("education")),
        bool(profile.get("skills")),
        bool(profile.get("experience"))
    ]

    completed_fields = sum(checks)
    total_fields = len(checks)

    return round(
        completed_fields / total_fields * 100
    )


def find_missing_information(
    profile: dict
) -> List[str]:
    """Find missing profile sections."""

    missing = []

    if not profile.get("full_name"):
        missing.append("Full name")

    if not profile.get("email"):
        missing.append("Email")

    if not profile.get("education"):
        missing.append("Education")

    if not profile.get("skills"):
        missing.append("Skills")

    if not profile.get("experience"):
        missing.append("Experience")

    return missing


# --------------------------------------------------
# Display candidate information
# --------------------------------------------------

def display_candidate_profile(
    profile: dict
) -> None:
    """Display the extracted candidate information."""

    completeness = calculate_completeness(
        profile
    )

    missing_information = (
        find_missing_information(profile)
    )

    st.subheader("Candidate Overview")

    metric1, metric2, metric3, metric4 = (
        st.columns(4)
    )

    metric1.metric(
        "Profile Completeness",
        f"{completeness}%"
    )

    metric2.metric(
        "Skills",
        len(profile.get("skills", []))
    )

    metric3.metric(
        "Education Records",
        len(profile.get("education", []))
    )

    metric4.metric(
        "Experience Records",
        len(profile.get("experience", []))
    )

    st.progress(
        completeness / 100
    )

    st.divider()

    information_column1, information_column2 = (
        st.columns(2)
    )

    with information_column1:
        st.text_input(
            "Full Name",
            value=profile.get(
                "full_name",
                ""
            ),
            disabled=True
        )

    with information_column2:
        st.text_input(
            "Email",
            value=profile.get(
                "email",
                ""
            ),
            disabled=True
        )

    st.divider()

    insights_tab, education_tab, skills_tab, experience_tab = (
        st.tabs(
            [
                "HR Insights",
                "Education",
                "Skills",
                "Experience"
            ]
        )
    )

    with insights_tab:
        st.subheader(
            "Missing Information"
        )

        if missing_information:
            for item in missing_information:
                st.warning(
                    f"Missing: {item}"
                )
        else:
            st.success(
                "All required profile sections were detected."
            )

        st.subheader(
            "Candidate Summary"
        )

        st.write(
            f"""
**Candidate:** {
    profile.get("full_name") or "Unknown"
}

**Email:** {
    profile.get("email") or "Not provided"
}

**Skills detected:** {
    len(profile.get("skills", []))
}

**Education records:** {
    len(profile.get("education", []))
}

**Experience records:** {
    len(profile.get("experience", []))
}
"""
        )

    with education_tab:
        education = profile.get(
            "education",
            []
        )

        if education:
            education_dataframe = (
                pd.DataFrame(education)
            )

            st.dataframe(
                education_dataframe,
                use_container_width=True,
                hide_index=True
            )
        else:
            st.info(
                "No education information was found."
            )

    with skills_tab:
        skills = profile.get(
            "skills",
            []
        )

        if skills:
            skill_columns = st.columns(4)

            for index, skill in enumerate(skills):
                column_number = index % 4

                with skill_columns[column_number]:
                    st.success(skill)
        else:
            st.info(
                "No skills were found."
            )

    with experience_tab:
        experience = profile.get(
            "experience",
            []
        )

        if experience:
            experience_dataframe = (
                pd.DataFrame(experience)
            )

            st.dataframe(
                experience_dataframe,
                use_container_width=True,
                hide_index=True
            )
        else:
            st.info(
                "No experience information was found."
            )


# --------------------------------------------------
# Streamlit interface
# --------------------------------------------------

st.title("👤 HR Candidate Profile Parser")

st.write(
    """
Upload a candidate resume or paste a resume snippet.

The application extracts the candidate's name, email,
education, skills and work experience into structured JSON.
"""
)

with st.sidebar:
    st.header("Project Information")

    st.write(
        """
**Technologies used**

- Streamlit
- LangChain JsonOutputParser
- Pydantic
- Hugging Face Inference API
- PyMuPDF
- Python DOCX
"""
    )

    st.warning(
        "AI-generated information should always be reviewed."
    )


uploaded_file = st.file_uploader(
    "Upload Candidate Resume",
    type=[
        "pdf",
        "docx",
        "txt"
    ]
)


manual_text = st.text_area(
    "Or paste a resume snippet",
    height=220,
    placeholder=(
        "Ahmed Ali\n"
        "ahmed@email.com\n\n"
        "Education:\n"
        "B.Sc. Computer Science, Cairo University, 2025\n\n"
        "Skills:\n"
        "Python, SQL, Machine Learning\n\n"
        "Experience:\n"
        "Software Intern at ABC Company, 2024"
    )
)


parse_button = st.button(
    "Parse Candidate Profile",
    type="primary",
    use_container_width=True
)


if parse_button:

    if uploaded_file is None and not manual_text.strip():
        st.error(
            "Please upload a resume or paste resume text."
        )

    else:
        try:
            with st.spinner(
                "Extracting and parsing candidate information..."
            ):

                if uploaded_file is not None:
                    resume_text = extract_resume_text(
                        uploaded_file
                    )
                else:
                    resume_text = (
                        manual_text.strip()
                    )

                if not resume_text:
                    raise ValueError(
                        "No readable text was found."
                    )

                candidate_profile = (
                    parse_candidate_profile(
                        resume_text
                    )
                )

                st.session_state[
                    "candidate_profile"
                ] = candidate_profile

                st.session_state[
                    "resume_text"
                ] = resume_text

            st.success(
                "Candidate profile parsed successfully."
            )

        except Exception as error:
            st.error(
                f"Error: {error}"
            )


if "candidate_profile" in st.session_state:

    profile = st.session_state[
        "candidate_profile"
    ]

    display_candidate_profile(
        profile
    )

    st.divider()

    json_output = json.dumps(
        profile,
        indent=4,
        ensure_ascii=False
    )

    with st.expander(
        "View Structured JSON"
    ):
        st.code(
            json_output,
            language="json"
        )

    with st.expander(
        "View Extracted Resume Text"
    ):
        st.text_area(
            "Extracted Text",
            value=st.session_state.get(
                "resume_text",
                ""
            ),
            height=300,
            disabled=True
        )

    st.download_button(
        label="Download Candidate JSON",
        data=json_output,
        file_name="candidate_profile.json",
        mime="application/json",
        use_container_width=True
    )